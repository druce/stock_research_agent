#!/usr/bin/env python3
"""
Final Report Assembly Phase

Combines all research outputs into polished final report with multi-format export.

Usage:
    ./skills/research_final.py SYMBOL [--work-dir DIR]

Output:
    - final_report.md
    - final_report.docx (if pandoc or python-docx available)
    - final_report.html (if pandoc or markdown available)
"""

import sys
import os
import argparse
import json
import logging
import subprocess
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import pandas as pd

# Template engine
from jinja2 import Environment, FileSystemLoader

# Import configuration
from config import (
    TEMPLATES_DIR,
    FINAL_REPORT_TEMPLATE,
    DATE_FORMAT_FILE,
)

# Import utilities
from utils import (
    setup_logging,
    validate_symbol,
    ensure_directory,
)

# Set up logging
logger = setup_logging(__name__)


def load_all_data(work_dir: Path, symbol: str) -> Dict[str, Any]:
    """
    Load all research data including deep research output.

    Extends the load_data function from research_report.py with additional
    deep research data.

    Args:
        work_dir: Work directory path
        symbol: Stock ticker symbol

    Returns:
        Dictionary containing all research data

    Example:
        >>> from pathlib import Path
        >>> data = load_all_data(Path('work/TSLA_20260116'), 'TSLA')
    """
    # Start with basic structure
    data: Dict[str, Any] = {
        'symbol': symbol,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'company_name': 'N/A',
        'sector': 'N/A',
        'industry': 'N/A',
        'latest_price': 'N/A',
        'market_cap': 'N/A',
        'revenue': 'N/A',
        'profit_margin': 'N/A',
        'roe': 'N/A',
        'trailing_pe': 'N/A',
    }

    # Load technical analysis
    tech_path = work_dir / '01_technical' / 'technical_analysis.json'
    if tech_path.exists():
        try:
            with tech_path.open('r') as f:
                data['technical_analysis'] = json.load(f)
                if 'latest_price' in data['technical_analysis']:
                    data['latest_price'] = f"{data['technical_analysis']['latest_price']:.2f}"
        except (IOError, json.JSONDecodeError) as e:
            logger.warning(f"Could not load technical analysis: {e}")

    # Load peers list with enhanced metrics
    peers_path = work_dir / '01_technical' / 'peers_list.json'
    ratios_path = work_dir / '02_fundamental' / 'key_ratios.csv'

    if peers_path.exists():
        try:
            with peers_path.open('r') as f:
                peers_data = json.load(f)
        except (IOError, json.JSONDecodeError) as e:
            logger.warning(f"Could not load peers data: {e}")
            peers_data = None

        # Load key ratios for peer metrics
        ratios_df = None
        if ratios_path.exists():
            try:
                ratios_df = pd.read_csv(ratios_path)
            except (IOError, pd.errors.ParserError) as e:
                logger.warning(f"Could not load ratios data: {e}")

        # Convert to list of dicts with enhanced metrics
        if peers_data and 'symbol' in peers_data and isinstance(peers_data['symbol'], list):
            peers: List[Dict[str, Any]] = []
            for i, peer_symbol in enumerate(peers_data['symbol']):
                peer_info: Dict[str, Any] = {
                    'symbol': peer_symbol,
                    'name': peers_data.get('name', [])[i] if i < len(peers_data.get('name', [])) else 'N/A',
                    'price': f"{peers_data.get('price', [])[i]:.2f}" if i < len(peers_data.get('price', [])) and peers_data.get('price', [])[i] else 'N/A',
                    'market_cap': f"{peers_data.get('market_cap', [])[i]:,.0f}" if i < len(peers_data.get('market_cap', [])) and peers_data.get('market_cap', [])[i] else 'N/A',
                }

                # Add financial metrics from ratios_df if available
                if ratios_df is not None and peer_symbol in ratios_df.columns:
                    # P/E Ratio
                    pe_row = ratios_df[ratios_df['Metric'] == 'Trailing P/E']
                    if not pe_row.empty:
                        pe_val = pe_row[peer_symbol].values[0]
                        peer_info['pe_ratio'] = f"{pe_val:.2f}" if pd.notna(pe_val) else 'N/A'
                    else:
                        peer_info['pe_ratio'] = 'N/A'

                    # Revenue
                    rev_row = ratios_df[ratios_df['Metric'] == 'Revenue (ttm)']
                    if not rev_row.empty:
                        rev_val = rev_row[peer_symbol].values[0]
                        if pd.notna(rev_val):
                            try:
                                peer_info['revenue'] = f"${float(rev_val)/1e9:.1f}B"
                            except (ValueError, TypeError) as e:
                                logger.debug(f"Could not format revenue for {peer_symbol}: {e}")
                                peer_info['revenue'] = 'N/A'
                        else:
                            peer_info['revenue'] = 'N/A'
                    else:
                        peer_info['revenue'] = 'N/A'

                    # Profit Margin
                    margin_row = ratios_df[ratios_df['Metric'] == 'Profit Margin']
                    if not margin_row.empty:
                        margin_val = margin_row[peer_symbol].values[0]
                        if pd.notna(margin_val):
                            try:
                                peer_info['profit_margin'] = f"{float(margin_val)*100:.1f}%"
                            except (ValueError, TypeError) as e:
                                logger.debug(f"Could not format margin for {peer_symbol}: {e}")
                                peer_info['profit_margin'] = 'N/A'
                        else:
                            peer_info['profit_margin'] = 'N/A'
                    else:
                        peer_info['profit_margin'] = 'N/A'

                    # ROE
                    roe_row = ratios_df[ratios_df['Metric'] == 'Return on Equity']
                    if not roe_row.empty:
                        roe_val = roe_row[peer_symbol].values[0]
                        if pd.notna(roe_val):
                            try:
                                peer_info['roe'] = f"{float(roe_val)*100:.1f}%"
                            except (ValueError, TypeError) as e:
                                logger.debug(f"Could not format ROE for {peer_symbol}: {e}")
                                peer_info['roe'] = 'N/A'
                        else:
                            peer_info['roe'] = 'N/A'
                    else:
                        peer_info['roe'] = 'N/A'
                else:
                    peer_info['pe_ratio'] = 'N/A'
                    peer_info['revenue'] = 'N/A'
                    peer_info['profit_margin'] = 'N/A'
                    peer_info['roe'] = 'N/A'

                peers.append(peer_info)

            data['peers'] = peers[:10]

    # Check for chart
    chart_path = work_dir / '01_technical' / 'chart.png'
    if chart_path.exists():
        data['chart_path'] = '01_technical/chart.png'

    # Check for income statement Sankey chart
    sankey_path = work_dir / '02_fundamental' / 'income_statement_sankey.png'
    if sankey_path.exists():
        data['income_statement_sankey_path'] = '02_fundamental/income_statement_sankey.png'

    # Load fundamental data
    overview_path = work_dir / '02_fundamental' / 'company_overview.json'
    if overview_path.exists():
        overview: Dict[str, Any] = {}
        try:
            with overview_path.open('r') as f:
                overview = json.load(f)
                data['company_name'] = overview.get('company_name', 'N/A')
                data['sector'] = overview.get('sector', 'N/A')
                data['industry'] = overview.get('industry', 'N/A')
        except (IOError, json.JSONDecodeError) as e:
            logger.warning(f"Could not load overview: {e}")
        if overview.get('market_cap') != 'N/A':
            data['market_cap'] = f"{overview['market_cap']:,.0f}"

        # Financial metrics
        data['revenue'] = f"{overview.get('revenue', 0):,.0f}" if overview.get('revenue') != 'N/A' else 'N/A'

        # Margins
        if overview.get('profit_margin') != 'N/A' and overview.get('profit_margin') is not None:
            data['profit_margin'] = f"{overview['profit_margin']*100:.2f}"
        else:
            data['profit_margin'] = 'N/A'

        # Returns
        if overview.get('roe') != 'N/A' and overview.get('roe') is not None:
            data['roe'] = f"{overview['roe']*100:.2f}"
        else:
            data['roe'] = 'N/A'

        # Valuation
        data['trailing_pe'] = f"{overview.get('trailing_pe', 0):.2f}" if overview.get('trailing_pe') != 'N/A' else 'N/A'

    # Load deep research output
    deep_output_path = work_dir / '08_deep_research' / 'deep_research_output.md'
    if deep_output_path.exists():
        try:
            with deep_output_path.open('r') as f:
                deep_content = f.read()
                data['deep_research_output'] = deep_content

                # Try to extract summary and conclusion sections
                # Look for "## 1" or "1." for summary
                summary_match = re.search(r'#+\s*1\.?\s*.*?[Ss]ummary.*?\n+(.*?)(?=\n#+|\Z)', deep_content, re.DOTALL)
                if summary_match:
                    data['deep_summary'] = summary_match.group(1).strip()

                # Look for "## 12" or "12." for conclusion
                conclusion_match = re.search(r'#+\s*12\.?\s*.*?[Cc]onclusion.*?\n+(.*?)(?=\n#+|\Z)', deep_content, re.DOTALL)
                if conclusion_match:
                    data['deep_conclusion'] = conclusion_match.group(1).strip()
        except IOError as e:
            logger.warning(f"Could not load deep research output: {e}")
            data['deep_research_output'] = '*Deep research not yet completed.*'
    else:
        data['deep_research_output'] = '*Deep research not yet completed.*'

    return data


def generate_final_report(data: Dict[str, Any], work_dir: Path) -> Path:
    """
    Generate final report using template.

    Args:
        data: Dictionary containing all research data
        work_dir: Work directory path

    Returns:
        Path to generated markdown report

    Example:
        >>> from pathlib import Path
        >>> report_path = generate_final_report(data, Path('work/TSLA_20260116'))
    """
    # Setup Jinja2 environment
    template_dir = Path(__file__).parent.parent / TEMPLATES_DIR
    env = Environment(loader=FileSystemLoader(str(template_dir)))

    # Custom filter for number formatting
    def format_number(value: Any) -> str:
        try:
            return f"{int(value):,}"
        except (ValueError, TypeError):
            return str(value)

    env.filters['format_number'] = format_number

    # Load template
    template = env.get_template(FINAL_REPORT_TEMPLATE)

    # Render template
    report_content = template.render(**data)
    report_content = report_content.replace('$', '\\$')

    # Save markdown file
    report_path = work_dir / 'final_report.md'
    with report_path.open('w') as f:
        f.write(report_content)

    logger.info(f"Saved: {report_path}")
    print(f"✓ Saved: {report_path}")
    return report_path


def convert_to_docx(md_path: Path, docx_path: Path) -> bool:
    """
    Convert markdown to Word document using pandoc or python-docx.

    Args:
        md_path: Path to markdown file
        docx_path: Path for output docx file

    Returns:
        True if conversion succeeded, False otherwise
    """
    try:
        # Try pandoc first (most reliable)
        result = subprocess.run(
            ['pandoc', str(md_path), '-o', str(docx_path)],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            logger.info(f"Saved: {docx_path} (via pandoc)")
            print(f"✓ Saved: {docx_path} (via pandoc)")
            return True
        else:
            logger.warning(f"Pandoc conversion failed: {result.stderr}")
            print(f"⚠ Pandoc conversion failed: {result.stderr}")

    except FileNotFoundError:
        logger.info("Pandoc not found")
        print("⚠ Pandoc not found. Install with: brew install pandoc")
    except subprocess.TimeoutExpired:
        logger.error("Pandoc timed out after 60 seconds")
        print(f"⚠ Pandoc timed out")
    except Exception as e:
        logger.error(f"Pandoc error: {e}", exc_info=True)
        print(f"⚠ Pandoc error: {e}")

    # Try python-docx as fallback
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Read markdown content
        with open(md_path, 'r') as f:
            content = f.read()

        # Simple markdown to docx conversion
        for line in content.split('\n'):
            if line.startswith('# '):
                # Heading 1
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # Heading 2
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Heading 3
                p = doc.add_heading(line[4:], level=3)
            elif line.strip() == '---':
                # Horizontal rule (skip)
                continue
            elif line.startswith('| '):
                # Table row (simplified handling)
                p = doc.add_paragraph(line)
                p.style = 'Normal'
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph(line.strip('**'))
                p.runs[0].bold = True
            else:
                # Regular paragraph
                if line.strip():
                    doc.add_paragraph(line)

        doc.save(str(docx_path))
        logger.info(f"Saved: {docx_path} (via python-docx)")
        print(f"✓ Saved: {docx_path} (via python-docx)")
        return True

    except ImportError:
        logger.info("python-docx not found")
        print("⚠ python-docx not found. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"python-docx error: {e}", exc_info=True)
        print(f"⚠ python-docx error: {e}")

    return False


def convert_to_html(md_path: Path, html_path: Path) -> bool:
    """
    Convert markdown to HTML using pandoc or markdown library.

    Args:
        md_path: Path to markdown file
        html_path: Path for output HTML file

    Returns:
        True if conversion succeeded, False otherwise
    """
    try:
        # Try pandoc first (most reliable)
        result = subprocess.run(
            ['pandoc', str(md_path), '-o', str(html_path), '--standalone', '--toc', '--css', 'style.css'],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            logger.info(f"Saved: {html_path} (via pandoc)")
            print(f"✓ Saved: {html_path} (via pandoc)")
            return True
        else:
            logger.warning(f"Pandoc conversion failed: {result.stderr}")
            print(f"⚠ Pandoc conversion failed: {result.stderr}")

    except FileNotFoundError:
        logger.info("Pandoc not found for HTML conversion")
        print("⚠ Pandoc not found for HTML conversion")
    except subprocess.TimeoutExpired:
        logger.error("Pandoc HTML conversion timed out")
        print(f"⚠ Pandoc timed out")
    except Exception as e:
        logger.error(f"Pandoc HTML error: {e}", exc_info=True)
        print(f"⚠ Pandoc error: {e}")

    # Try markdown library as fallback
    try:
        import markdown

        with md_path.open('r') as f:
            md_content = f.read()

        # Convert markdown to HTML with extensions
        html_content = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code', 'toc', 'meta']
        )

        # Wrap in basic HTML template with styling
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Equity Research Report</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        h1 {{
            color: #333;
            border-bottom: 2px solid #333;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #555;
            margin-top: 30px;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 30px 0;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

        with html_path.open('w') as f:
            f.write(full_html)

        logger.info(f"Saved: {html_path} (via markdown library)")
        print(f"✓ Saved: {html_path} (via markdown library)")
        return True

    except ImportError:
        logger.info("markdown library not found")
        print("⚠ markdown library not found. Install with: pip install markdown")
    except (IOError, Exception) as e:
        logger.error(f"markdown library error: {e}", exc_info=True)
        print(f"⚠ markdown library error: {e}")

    return False


def main() -> int:
    """
    Main execution function.

    Returns:
        Exit code (0 for success, 1 for failure)
    """
    parser = argparse.ArgumentParser(description='Final report assembly phase')
    parser.add_argument('symbol', help='Stock ticker symbol')
    parser.add_argument(
        '--work-dir',
        default=None,
        help='Work directory path (default: work/SYMBOL_YYYYMMDD)'
    )
    args = parser.parse_args()

    # Normalize symbol
    symbol = validate_symbol(args.symbol)

    # Generate work directory if not specified
    if not args.work_dir:
        date_str = datetime.now().strftime(DATE_FORMAT_FILE)
        work_dir = Path('work') / f'{symbol}_{date_str}'
    else:
        work_dir = Path(args.work_dir)

    print("="*60)
    print("Final Report Assembly Phase")
    print("="*60)
    print(f"Symbol: {symbol}")
    print(f"Work Directory: {work_dir}")
    print("="*60)

    try:
        # Load all data
        logger.info("Loading all research data...")
        print(f"\nLoading all research data...")
        data = load_all_data(work_dir, symbol)
        logger.info("Data loaded")
        print(f"✓ Data loaded")

        # Generate final report
        logger.info("Generating final report...")
        print(f"\nGenerating final report...")
        md_path = generate_final_report(data, work_dir)

        # Convert to other formats
        logger.info("Converting to additional formats...")
        print(f"\nConverting to additional formats...")

        # DOCX
        docx_path = work_dir / 'final_report.docx'
        convert_to_docx(md_path, docx_path)

        # HTML
        html_path = work_dir / 'final_report.html'
        convert_to_html(md_path, html_path)

        print("\n" + "="*60)
        print("SUCCESS: Final report assembly completed!")
        print("="*60)
        print(f"\nOutputs:")
        print(f"  - {md_path}")
        if docx_path.exists():
            print(f"  - {docx_path}")
        if html_path.exists():
            print(f"  - {html_path}")

        logger.info("Final report assembly completed")
        return 0

    except (IOError, KeyError) as e:
        logger.error(f"Error in final report assembly: {e}", exc_info=True)
        print(f"\n❌ Error in final report assembly: {e}")
        return 1
    except Exception as e:
        logger.error(f"Unexpected error in final report assembly: {e}", exc_info=True)
        print(f"\n❌ Error in final report assembly: {e}")
        return 1


if __name__ == '__main__':
    sys.exit(main())
